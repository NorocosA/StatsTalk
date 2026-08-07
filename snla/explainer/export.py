"""Word report export for StatsTalk analysis results."""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any

from snla.parser.schema import AnalysisResult, TableResult, dict_to_analysis_result


def export_to_docx(
    output_path: str,
    user_query: str,
    method: str,
    analysis_result: AnalysisResult | dict[str, Any],
    explanation: str,
    data_file: str = "",
    export_apa: bool = True,
    parameters: dict[str, Any] | None = None,
) -> str:
    """Generate a Word report and return the absolute output path."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    result = _coerce_result(analysis_result)
    stats = result.statistics or {}

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)

    title = doc.add_heading("StatsTalk 统计分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if data_file:
        doc.add_paragraph(f"数据文件: {os.path.basename(data_file)}")

    doc.add_heading("1. 分析问题", level=1)
    doc.add_paragraph(user_query or "未记录原始问题")

    doc.add_heading("2. 统计方法", level=1)
    doc.add_paragraph(f"推荐方法: {_method_label(method)}")
    if parameters and parameters.get("alpha") is not None:
        doc.add_paragraph(f"显著性水平: α = {parameters['alpha']:g}")
    if result.n_valid:
        doc.add_paragraph(f"有效样本量: N = {result.n_valid}")

    doc.add_heading("3. 关键统计结果", level=1)
    if stats:
        table = doc.add_table(rows=1, cols=2, style="Table Grid")
        table.rows[0].cells[0].text = "统计量"
        table.rows[0].cells[1].text = "数值"
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for label, keys, fmt in _STAT_ROWS:
            value = _first_present(stats, keys)
            _add_stat_row(table, label, value, fmt)
    else:
        doc.add_paragraph("本次结果未提取到可汇总的关键统计量。")

    if result.tables:
        doc.add_heading("4. 结果表", level=1)
        for table_result in result.tables[:3]:
            _add_result_table(doc, table_result)

    doc.add_heading("5. 结果解读", level=1)
    doc.add_paragraph(explanation or "暂无结果解读。")

    if export_apa:
        doc.add_heading("6. APA 摘要", level=1)
        paragraph = doc.add_paragraph(_build_apa(method, _method_label(method), stats))
        if paragraph.runs:
            paragraph.runs[0].bold = True

    doc.add_paragraph("")
    footer = doc.add_paragraph("本报告由 StatsTalk 自动生成。")
    if footer.runs:
        footer.runs[0].italic = True

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


_STAT_ROWS: tuple[tuple[str, tuple[str, ...], str], ...] = (
    ("p 值", ("p_value", "p", "p_val"), ".4f"),
    ("t 值", ("t_value", "t"), ".3f"),
    ("F 值", ("f_value", "f", "F"), ".3f"),
    ("卡方值", ("chi_square", "chi2"), ".3f"),
    ("自由度", ("df", "dof"), ""),
    ("效应量", ("cohen_d", "d_value", "eta_sq", "np2", "rbc", "effect_size"), ".3f"),
    ("R 平方", ("r_squared", "r2"), ".3f"),
    ("相关系数 r", ("r",), ".3f"),
    ("均值差", ("mean_diff",), ".2f"),
    ("样本量 N", ("n_valid", "n"), ""),
    ("均值", ("mean",), ".2f"),
    ("标准差", ("std_dev", "std_deviation", "sd"), ".2f"),
    ("最小值", ("minimum", "min"), ".2f"),
    ("最大值", ("maximum", "max"), ".2f"),
)


def _coerce_result(result: AnalysisResult | dict[str, Any]) -> AnalysisResult:
    if isinstance(result, AnalysisResult):
        return result
    if isinstance(result, dict):
        return dict_to_analysis_result(result)
    return AnalysisResult(notes=["Export received an unsupported result object."])


def _first_present(stats: dict[str, Any], keys: tuple[str, ...]) -> Any:
    for key in keys:
        value = stats.get(key)
        if value is not None:
            return value
    return None


def _add_stat_row(table: Any, label: str, value: Any, fmt: str = "") -> None:
    if value is None:
        return
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = _format_value(value, fmt)


def _add_result_table(doc: Any, table_result: TableResult) -> None:
    doc.add_paragraph(table_result.title or "结果表")
    rows = table_result.rows[:8]
    if not rows:
        doc.add_paragraph("该表没有可展示的行。")
        return

    columns: list[str] = []
    for row in rows:
        for key in row:
            if key not in columns:
                columns.append(key)
    columns = columns[:6]

    table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
    for idx, column in enumerate(columns):
        table.rows[0].cells[idx].text = column or "项目"
    for source_row in rows:
        row = table.add_row()
        for idx, column in enumerate(columns):
            row.cells[idx].text = str(source_row.get(column, ""))


def _format_value(value: Any, fmt: str = "") -> str:
    if isinstance(value, float) and fmt:
        try:
            return format(value, fmt)
        except (TypeError, ValueError):
            return str(value)
    return str(value)


def _method_label(method: str) -> str:
    labels = {
        "independent_t_test": "独立样本 t 检验",
        "paired_t_test": "配对样本 t 检验",
        "oneway_anova": "单因素方差分析",
        "pearson_correlation": "Pearson 相关分析",
        "spearman_correlation": "Spearman 秩相关分析",
        "correlations": "相关分析",
        "simple_regression": "简单线性回归",
        "multiple_regression": "多元线性回归",
        "logistic_regression": "Logistic 回归",
        "chi_square": "卡方检验",
        "crosstabs": "交叉表/卡方检验",
        "descriptives": "描述性统计",
        "frequencies": "频率分析",
        "mann_whitney_u": "Mann-Whitney U 检验",
        "kruskal_wallis": "Kruskal-Wallis 检验",
        "wilcoxon": "Wilcoxon 符号秩检验",
    }
    return labels.get(method, method.replace("_", " ").title())


def _build_apa(method: str, method_label: str, stats: dict[str, Any]) -> str:
    p = stats.get("p_value")
    if p is None:
        return "统计结果摘要: 关键显著性指标不足，建议参考上方结果表。"

    t_value = stats.get("t_value")
    df = stats.get("df")
    if t_value is not None and df is not None:
        return f"采用{method_label}，结果显示 t({int(df)}) = {t_value:.2f}, {_apa_p(p)}。"

    f_value = stats.get("f_value")
    if f_value is not None:
        df_text = f"({int(df)}, ...)" if df is not None else ""
        return f"采用{method_label}，结果显示 F{df_text} = {f_value:.2f}, {_apa_p(p)}。"

    r = stats.get("r")
    if r is not None:
        n = stats.get("n_valid") or stats.get("n")
        n_text = f", N = {int(n)}" if n is not None else ""
        return f"采用{method_label}，结果显示 r = {r:.3f}{n_text}, {_apa_p(p)}。"

    chi2 = stats.get("chi_square")
    if chi2 is not None:
        df_text = f"({int(df)})" if df is not None else ""
        return f"采用{method_label}，结果显示 chi-square{df_text} = {chi2:.2f}, {_apa_p(p)}。"

    u_value = stats.get("u") or stats.get("u_value")
    if u_value is not None:
        return f"采用{method_label}，结果显示 U = {u_value:.2f}, {_apa_p(p)}。"

    h_value = stats.get("h") or stats.get("h_value")
    if h_value is not None:
        return f"采用{method_label}，结果显示 H = {h_value:.2f}, {_apa_p(p)}。"

    return f"采用{method_label}，结果显示 {_apa_p(p)}。"


def _apa_p(p_value: float) -> str:
    if p_value < 0.001:
        return "p < .001"
    return f"p = {p_value:.3f}"
