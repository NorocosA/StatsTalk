from __future__ import annotations

import asyncio
import base64
import json

from docx import Document

from snla.analysis import AnalysisAudit, AnalysisSuccess
from snla.explainer.export import export_to_docx
from snla.explainer.record import build_analysis_record
from snla.parser.schema import AnalysisResult, TableResult


def _audit(**overrides):
    values = {
        "request_id": "record-123",
        "started_at": "2026-08-11T10:00:00+00:00",
        "completed_at": "2026-08-11T10:00:01+00:00",
        "method": "independent_t_test",
        "preferred_backend": "spss",
        "effective_backend": "python",
        "parser_used": "python_pingouin",
        "fallback_reason": {"code": "SPSS_EXECUTABLE_NOT_FOUND"},
    }
    values.update(overrides)
    return AnalysisAudit(**values)


def _result():
    return AnalysisResult(
        analysis_type="T-TEST",
        tables=[
            TableResult(
                title="Independent Samples Test",
                rows=[{"t": 2.34, "df": 18, "p": 0.021}],
                notes=["Equal variances were not assumed."],
            )
        ],
        statistics={"t_value": 2.34, "df": 18, "p_value": 0.021},
        n_valid=20,
        n_missing=2,
        notes=["Small samples can make the estimate unstable."],
        raw_output_path=r"C:\Users\student\private-output.xml",
        parser_used="python_pingouin",
    )


def test_success_payload_has_summary_and_complete_advanced_layers():
    fallback = {
        "code": "SPSS_EXECUTABLE_NOT_FOUND",
        "message": "SPSS was unavailable; Python was used.",
        "method": "independent_t_test",
        "announce": True,
    }
    outcome = AnalysisSuccess(
        user_query="compare groups",
        method="independent_t_test",
        backend="python",
        plan_explanation="user selected",
        result=_result(),
        explanation="The groups differ significantly.",
        syntax="T-TEST GROUPS=group(1 2) /VARIABLES=score.",
        warning="Python validation warning.",
        parameters={"alpha": 0.05},
        variable_roles={"grouping_variable": "group", "test_variable": "score"},
        fallback_reason=fallback,
        audit=_audit(fallback_reason=fallback),
    )

    payload = outcome.to_payload()
    summary = payload["summary"]
    advanced = payload["advanced"]

    assert summary["conclusion"] == "The groups differ significantly."
    assert summary["key_statistics"]["p_value"] == 0.021
    assert "Python validation warning." in summary["warnings"]
    assert "Small samples can make the estimate unstable." in summary["warnings"]
    assert "SPSS was unavailable; Python was used." in summary["warnings"]
    assert summary["fallback"]["used"] is True
    assert advanced["tables"][0]["rows"][0]["t"] == 2.34
    assert advanced["syntax"].startswith("T-TEST")
    assert advanced["backend"]["effective"] == "python"
    assert advanced["diagnostics"]["parser_used"] == "python_pingouin"
    assert payload["analysis_record"]["variables"] == [
        {"role": "grouping_variable", "name": "group"},
        {"role": "test_variable", "name": "score"},
    ]


def test_json_record_is_allowlisted_and_removes_private_fields():
    result = _result()
    result.statistics.update(
        {
            "access_token": "token-secret",
            "api_key": "sk-secret",
            "original_file_path": r"C:\Users\student\private.sav",
        }
    )
    result.tables[0].rows.append(
        {"mean": 4.2, "password": "secret-password", "raw_data": [1, 2, 3]}
    )

    record = build_analysis_record(
        capability="independent_t_test",
        variable_roles={"grouping_variable": "group", "test_variable": "score"},
        parameters={
            "alpha": 0.05,
            "api_key": "sk-param-secret",
            "source_path": r"C:\Users\student\private.sav",
        },
        result=result,
        conclusion="Conclusion",
        syntax="T-TEST.",
        preferred_backend="spss",
        effective_backend="python",
        fallback_reason={"code": "SPSS_EXECUTABLE_NOT_FOUND"},
        warning=None,
        greylist_warnings=(),
        audit=_audit(),
        selection_source="user_selection",
    )

    serialized = json.dumps(record)
    for forbidden in (
        "token-secret",
        "sk-secret",
        "sk-param-secret",
        "secret-password",
        "private.sav",
        "private-output.xml",
        "raw_data",
    ):
        assert forbidden not in serialized
    assert record["parameters"] == {"alpha": 0.05}
    assert record["advanced"]["tables"][0]["rows"][1] == {"mean": 4.2}


def test_word_export_contains_reproducibility_and_advanced_evidence(tmp_path):
    record = build_analysis_record(
        capability="independent_t_test",
        variable_roles={"grouping_variable": "group", "test_variable": "score"},
        parameters={"alpha": 0.05},
        result=_result(),
        conclusion="The groups differ significantly.",
        syntax="T-TEST GROUPS=group(1 2) /VARIABLES=score.",
        preferred_backend="spss",
        effective_backend="python",
        fallback_reason={
            "code": "SPSS_EXECUTABLE_NOT_FOUND",
            "message": "SPSS was unavailable; Python was used.",
        },
        warning="Interpret cautiously.",
        greylist_warnings=(),
        audit=_audit(),
        selection_source="user_selection",
    )
    output = tmp_path / "record.docx"

    export_to_docx(
        output_path=str(output),
        user_query="compare groups",
        method="independent_t_test",
        analysis_result=_result(),
        explanation="The groups differ significantly.",
        analysis_record=record,
    )

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Interpret cautiously." in text
    assert "SPSS was unavailable; Python was used." in text
    assert "grouping_variable: group" in text
    assert "实际执行后端: Python" in text
    assert "T-TEST GROUPS=group" in text
    assert "Independent Samples Test" in text
    assert "2.34" in table_text
    assert record["versions"]["statstalk"] in text


def test_mcp_json_export_returns_the_same_machine_readable_record(tmp_path, monkeypatch):
    from snla import mcp_server

    class Context:
        session_id = "record-export"

    record = {"schema_version": "1.0", "capability": "descriptives"}
    monkeypatch.setattr(mcp_server, "_upload_dir", tmp_path)
    mcp_server._session_states.clear()
    mcp_server._session_states[Context.session_id] = mcp_server.MCPState(
        last_result=_result(),
        last_record=record,
    )

    payload = asyncio.run(mcp_server.snla_export(Context(), format="json"))

    assert payload["ok"] is True
    assert payload["record"] == record
    assert json.loads(base64.b64decode(payload["content_base64"])) == record
